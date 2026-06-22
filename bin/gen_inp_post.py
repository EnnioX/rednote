# -*- coding: utf-8 -*-
"""Generate 磷化铟 InP publish post Word document."""
from docx import Document
from docx.shared import Pt, Emu, RGBColor

doc = Document()

for section in doc.sections:
    section.page_width = Emu(7772400)
    section.page_height = Emu(10058400)
    section.top_margin = Emu(899795)
    section.bottom_margin = Emu(899795)
    section.left_margin = Emu(899795)
    section.right_margin = Emu(899795)

style = doc.styles['Normal']
style.font.name = '微软雅黑'  # 微软雅黑
style.font.size = Pt(11)


def add_para(text, size_pt=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size_pt)
    run.bold = bold
    return p


def add_empty():
    doc.add_paragraph()


# === Meta ===
add_para(
    "post_id: POST-20260622-001 | post_date: 2026-06-22\n"
    "source: reports/20260622/磷化铟深挖.md\n"
    "topic: 磷化铟 | tags: 磷化铟, InP, "
    "光芯片, 光模块, AI算力, "
    "铟, 化合物半导体 | image_count: 5",
    size_pt=9
)

add_empty()
add_empty()

# === Title ===
add_para('磷化铟，供给缺口超过70%', size_pt=18, bold=True)

add_empty()

# === Body ===

# P1 - Opening
add_para(
    "AI数据中心的每一次信号传输都依赖激光器。"
    "制造激光器芯片的磷化铟（InP）衬底，"
    "全球有效产能仅约75万片，"
    "而2026年需求已达260—300万片，"
    "缺口超过70%。"
    "三寡头订单排至2028年，"
    "2英寸衬底价格从800美元飙升至2,300—2,500美元。"
    "这是正在发生的供给端事实，不是预测。"
)
add_empty()

add_para("三个维度理解这个缺口的性质：")
add_empty()

# P2 - Demand
add_para(
    "第一，需求端发生了结构性量纲切换。"
    "光模块从800G到1.6T到3.2T，需求计量单位"
    "已从“按模块数”切换为“按通道数”"
    "——1.6T模块的InP消耗量是800G的2.7—2.8倍。"
    "英伟达预测2026—2030年InP晶圆需求激增约20倍，"
    "Lumentum给出的AI数据中心InP需求CAGR是85%。"
    "每一个英伟达Rubin平台交换机搭载18个硅光引擎，"
    "每个引擎都需要InP基CW激光器。"
    "光模块的数量和每个模块的通道数同时在增长，"
    "这是乘数效应而非线性增长。"
)
add_empty()

# P3 - Supply
add_para(
    "第二，供给端三重约束叠加："
    "寡头垄断+扩产周期长+铟出口管制。"
    "住友电工（42%）、AXT（36%）、JX金属（13%）"
    "三家控制全球超过90%的InP衬底产能。"
    "扩产瓶颈清晰——MOCVD设备交付7—12个月，"
    "良率爬坡8—12个月，客户认证12—24个月，"
    "从宣布扩产到形成有效产能至少2年。"
    "中国控制全球75%铟储量，"
    "2026年将铟出口锁定在年产量的30%以内，"
    "进一步锁紧供给。"
    "铟金属价格从2025年初约2,800元/kg涨至4,950元/kg（+76%）。"
)
add_empty()

# P4 - InP irreplaceable
add_para(
    "第三，也是最关键的一点："
    "在AI光互联中，磷化铟在光源环节不可替代。"
    "硅是间接带隙半导体，不能发光"
    "——这是物理定律，不是技术路线选择。"
    "CPO（共封装光学）和硅光技术"
    "可以替代部分调制和集成功能，"
    "但光源环节必须使用InP基激光器。"
    "薄膜铌酸锂替代的是调制器，不是光源。"
    "三材料的分工已明确：磷化铟做光源（发动机），"
    "硅光做无源集成（车身），"
    "薄膜铌酸锂做超高速调制（变速箱）。"
    "只要AI数据中心需要光信号，就需要磷化铟。"
)
add_empty()

# P5 - A-share mapping defects
add_para(
    "但需要如实指出："
    "A股映射存在显著的结构性缺陷。"
    "最核心的衬底标的云南锗业，"
    "磷化铟收入占比仅约13%，"
    "2025全年净利仅2,014万元，"
    "而PE已超过3,000倍。"
    "公司已于6月15日自行发布异动公告，"
    "明确警示"
    "“市盈率和市净率显著高于同行业平均水平，"
    "存在市场情绪过热、非理性炒作风险”。"
    "光芯片环节的源杰科技（营收约3亿）、"
    "仕佳光子（有源芯片占比不到15%），"
    "体量小且100G以上产品仍处于客户验证期。"
    "当前A股没有一家以磷化铟衬底或光芯片"
    "为主营收入超过50%的上市公司。"
    "产业逻辑成立，但映射标的的盈利质量和估值水平"
    "与产业景气之间存在断崖式落差。"
)
add_empty()

# P6 - Core judgment
add_para(
    "这也是深挖报告给出的核心判断："
    "磷化铟的产业供需缺口真实，"
    "光源环节的不可替代性有物理定律支撑，"
    "但从A股投资角度看，"
    "核心标的的PE/业绩匹配度极差，"
    "产业逻辑的兑现周期（3—5年）远超策略持有周期。"
    "整条方向建议作为“产业前沿跟踪”关注，"
    "而非进入观察名单。"
)
add_empty()

# P7 - Next steps
add_para(
    "下一步验证指标："
    "7—8月中报——云南锗业化合物半导体"
    "营收增速是否超过50%、"
    "三安光电光技术板块是否首次单独披露营收；"
    "200G EML国产光芯片能否在2026下半年通过客户认证；"
    "11月铟出口管制暂缓到期后的政策走向。"
    "失效条件方面，最致命的信号是硅基光源"
    "（锗硅/量子点激光器）取得工程突破"
    "——若硅能发光，InP的不可替代性将被物理层面动摇。"
    "但目前此类技术仍处于实验室阶段，"
    "距产业化至少3—5年。"
    "短期内更现实的失效条件是："
    "铟出口管制放宽、全球InP产能超预期释放、"
    "或AI资本开支全局性缩减导致光模块出货量增速降至20%以下。"
)
add_empty()

# === Information & Data Sources ===
add_para("Information & Data Sources：")

evidence_lines = [
    "[1] 2026-06｜Yole/LightCounting/野村｜市场报告",
    "[2] 2026-06｜Lumentum/英伟达｜公司公告",
    "[3] 2026-06｜华泰证券/行业共识｜行业调研",
    "[4] 2026-06｜LightCounting/行业分析｜市场报告",
    "[5] 2026-06｜住友电工/AXT/JX金属财报｜公司公告",
    "[6] 2026-06｜AXT/Coherent/JX金属公告｜公司公告",
    "[7] 2026-06｜英伟达/Coherent｜公司公告",
    "[8] 2026-06｜中信建投/国金证券｜行业调研",
    "[9] 2026-06-15｜云南锗业公告｜公司公告",
    "[10] 2026｜商务部/外交部｜政府公告",
    "[11] 2026-06-12｜东方财富/证券时报｜市场报告",
]

for line in evidence_lines:
    add_para(line, size_pt=10)

add_empty()

# === Tags ===
add_para(
    "#磷化铟 #InP #光芯片 #光模块 "
    "#AI算力 #铟 #化合物半导体 #A股",
    size_pt=10
)

add_empty()

# === Disclaimer ===
add_para(
    "⚠️ 本内容仅供信息分享与逻辑交流，"
    "不构成任何投资建议。"
    "股市有风险，决策需谨慎。",
    size_pt=9
)

add_empty()

# === Slogan ===
add_para(
    "基于当下产业事实与公开数据，"
    "以图表提炼关键变化，"
    "不作具体投资判断。",
    size_pt=9
)

# Save
output_path = (
    "/Users/ennio/Documents/AI_work/rednote/posts/20260622-1/"
    "20-1-磷化铟缺口超过70%.docx"
)
doc.save(output_path)
print("Saved:", output_path)
