"""生成船舶制造超级周期发布稿 .docx"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# 元数据（正文第一页首行，三段连续）
meta_lines = [
    'post_id: POST-20260622-001 | post_date: 2026-06-22',
    'source: reports/20260622/船舶制造超级周期深挖.md',
    'topic: 船舶制造超级周期 | tags: 船舶制造, 造船周期, 中国船舶, 中国动力, 中船防务, 全球贸易 | image_count: 5',
]
for line in meta_lines:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# 空行分隔
p = doc.add_paragraph('')

# === 正文 ===

# 封面标题（开头点出变化）
p = doc.add_paragraph('造船超级周期，利润开始兑现')
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

p = doc.add_paragraph('')

# 开头段：核心事实
p = doc.add_paragraph(
    '船舶制造板块Q1出现了三份需要认真对待的财报。中国船舶归母净利+252%至48亿，'
    '中国动力+49%至5.88亿，中船防务+115%至3.96亿。三家核心标的的利润加速释放'
    '不是预测，是已兑现的数据。手持订单4,674亿排至2029-2030年，未来三年的收入已在合同里。'
)

p = doc.add_paragraph('')

# 中间：三个结构性事实
p = doc.add_paragraph('三个独立于宏观经济波动的结构性事实。')

p = doc.add_paragraph('')

p = doc.add_paragraph(
    '第一，船队老龄化——物理折旧决定了替换底线。全球商船队45%船龄超过15年，'
    '约33%现有船舶将在未来十年面临更新。即便全球贸易零增长，存量替换就撑起'
    '约5,000万载重吨/年的新船订单。这不是需求预测，是资产折旧的物理节奏。'
)

p = doc.add_paragraph('')

p = doc.add_paragraph(
    '第二，产能硬约束——供给恢复追不上需求释放。全球活跃船厂较2010年峰值下降35%，'
    '中国从400余家缩至100-130家。熟练焊工培养3-5年，大型船坞建设周期3-5年。'
    '即使扩产项目全部落地，产能缺口也将维持到2028年以后。这种供需错配的持续时间，'
    '是本轮周期与历史上任何一轮最本质的区别。'
)

p = doc.add_paragraph('')

p = doc.add_paragraph(
    '第三，新船价格高位运行——锁定未来盈利。克拉克森新船价格指数182点，'
    '处于历史95%分位，较2020年低位上涨47%。前期2023-2024年高价订单已开始进入交付：'
    '中国船舶Q1毛利率17.48%，同比+4.57pct；净利率11.94%，同比+5.51pct。'
    'Q2起高价订单交付占比将进一步提升，毛利率有望向20%+突破。'
)

p = doc.add_paragraph('')

# A股映射
p = doc.add_paragraph(
    'A股映射完整且纯正度高。整船总装——中国船舶（造船收入>90%，全球最大造船集团），'
    '船用动力——中国动力（全球船用发动机份额>60%，排产至2028-2029），'
    '军民船——中船防务（军品保底+民品弹性，Q1营收+56%），'
    '锚链——亚星锚链（全球锚链份额>50%）。'
    '四个标的的核心收入均可明确追溯至船舶业务，产业链从总装到发动机到关键部件形成闭环。'
    '这是当前A股唯一与AI Capex零相关的独立产业方向——有独立的景气逻辑和业绩兑现。'
)

p = doc.add_paragraph('')

# 风险
p = doc.add_paragraph('三件需要诚实面对的事。')

p = doc.add_paragraph('')

p = doc.add_paragraph(
    '周期位置：造船上行周期已行至中段偏后，而非早期。2025年全球新签订单同比-27%，'
    '增速已从"高速"转入"平稳"。克拉克森指数182点处于95%分位——历史上接近该分位后'
    '均出现不同程度回调。当前值得辩论的问题不是"周期是否上行"，而是"上行还剩多少空间"。'
)

p = doc.add_paragraph('')

p = doc.add_paragraph(
    '交付潮与贸易风险：2027-2028年全球交付量增速将超20%。若集中交付叠加贸易需求'
    '疲弱——WTO已将2026年全球货物贸易增速大幅下调至0.5%——可能形成"交付潮→运价跌→'
    '订单减"的负反馈，反作用于新船订单。'
)

p = doc.add_paragraph('')

p = doc.add_paragraph(
    '政策悬剑：美国301调查虽暂停一年但未解除，港口费预期始终存在。'
    '若2026H2恢复调查并加征费用，中国新接订单份额可能从84.9%回落至65-70%，'
    '对板块估值溢价形成压制。'
)

p = doc.add_paragraph('')

# 结尾：验证指标与判断
p = doc.add_paragraph(
    '下一步看8月中报。中国船舶Q2毛利率若突破20%（Q1为17.48%），'
    '将验证高价订单交付拐点；中国动力燃气轮机订单增速能否维持>50%；'
    '中船防务Q2营收是否逼近全年目标212亿的一半。克拉克森指数若连续跌破178，'
    '需触发复核。失效条件：全球贸易量连续两季下降>3%，或中国船舶毛利率连续两季环比下降。'
)

p = doc.add_paragraph('')

p = doc.add_paragraph(
    '综合来看：船舶方向的基本面确认度高（Q1三份财报交叉验证），估值合理'
    '（龙头2026E PE 14-20x），拥挤度远低于AI/科技等赛道。周期位置决定了追高的性价比'
    '在降低，但作为组合中独立于AI产业逻辑的防御性配置，有其独特价值。'
)

p = doc.add_paragraph('')

# === Information & Data Sources ===
p = doc.add_paragraph('Information & Data Sources：')

evidence_items = [
    '2026-04｜中国船舶2026Q1财报｜公司公告',
    '2026-06｜克拉克森/东吴证券/招商证券｜行业报告',
    '2026-06｜中国动力/中船防务Q1财报｜公司公告',
    '2026-06｜克拉克森｜行业报告',
    '2026-06｜招商证券年度策略｜行业报告',
    '2026-06｜招商证券/克拉克森｜行业报告',
    '2026-06｜BIMCO/克拉克森｜行业报告',
    '2026-06｜克拉克森/行业报道｜行业报告',
    '2026-06｜WTO｜政府公告',
    '2026-06｜USTR/美国国会｜政府公告',
]

for i, item in enumerate(evidence_items, 1):
    p = doc.add_paragraph(f'[{i}] {item}')

p = doc.add_paragraph('')

# 标签
p = doc.add_paragraph('#A股 #船舶制造 #造船周期 #中国船舶 #中国动力 #全球贸易 #资本市场 #周期股')

p = doc.add_paragraph('')

# 免责声明
p = doc.add_paragraph('⚠️ 本内容仅供信息分享与逻辑交流，不构成任何投资建议。股市有风险，决策需谨慎。')

# 保存
output_dir = '/Users/ennio/Documents/AI_work/rednote/posts/20260622-1'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, '20-1-船舶制造超级周期.docx')
doc.save(output_path)
print(f'发布稿已保存: {output_path}')
